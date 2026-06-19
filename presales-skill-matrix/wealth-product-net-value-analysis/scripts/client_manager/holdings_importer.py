"""
持仓信息导入引擎 v1.0
支持: 截图OCR识别 | Word/PDF导入 | 浏览器爬取 | Excel/CSV导出
"""
import json
import os
import re
import hashlib
from datetime import datetime
from pathlib import Path
from typing import Optional

SCRIPT_DIR = Path(__file__).resolve().parent
DATA_DIR = SCRIPT_DIR.parent.parent / "data"
CLIENTS_DIR = DATA_DIR / "clients"


class HoldingsImporter:
    """客户持仓信息导入引擎"""

    # 基金代码模式: 6位数字
    FUND_CODE_PATTERN = re.compile(r'(?<!\d)(\d{6})(?!\d)')

    # 份额模式
    SHARES_PATTERN = re.compile(
        r'(?:份额|持有|买入|份数)[:：\s]*(\d+(?:,\d{3})*(?:\.\d+)?)\s*(?:份|万份)?'
        r'|(\d+(?:,\d{3})*(?:\.\d+)?)\s*(?:份|万份)'
    )

    # 成本/净值模式
    COST_PATTERN = re.compile(
        r'(?:成本|买入价|净值|价格)[:：\s]*(\d+\.?\d*)'
    )

    # 基金名称模式
    FUND_NAME_PATTERN = re.compile(
        r'(?:基金名称|产品名称|名称)[:：\s]*([^\n\r]{2,30}?)(?:\s+(?:代码|份额|成本|买入|净值|日期|$))'
    )

    # 日期模式
    DATE_PATTERN = re.compile(
        r'(\d{4})[-/年](\d{1,2})[-/月](\d{1,2})[日号]?'
    )

    # 金额模式
    AMOUNT_PATTERN = re.compile(
        r'(?:金额|市值|投入)[:：\s]*(\d+(?:,\d{3})*(?:\.\d+)?)\s*(?:元|万元|万)?'
    )

    # 基金公司网站URL模式
    FUND_URL_PATTERNS = [
        (r'天天基金', 'eastmoney'),
        (r'fund\.eastmoney\.com', 'eastmoney'),
        (r'antfortune|蚂蚁|支付宝', 'antfortune'),
        (r'howbuy|好买', 'howbuy'),
        (r'lufax|陆金所', 'lufax'),
        (r'([a-zA-Z0-9_-]+)\.eastmoney\.com', 'eastmoney'),
    ]

    def __init__(self, data_dir=None):
        self.data_dir = data_dir or DATA_DIR
        self.clients_dir = self.data_dir / "clients"
        self.clients_dir.mkdir(parents=True, exist_ok=True)

    # ==================== 截图OCR识别 ====================

    def import_from_screenshot(self, image_path: str, client_id: str = None) -> dict:
        """
        从截图OCR识别持仓信息

        Args:
            image_path: 截图文件路径
            client_id: 客户ID（可选，用于保存到仓库）

        Returns:
            dict: {'success': bool, 'holdings': list, 'raw_text': str, 'errors': list}
        """
        result = {'success': False, 'holdings': [], 'raw_text': '', 'errors': []}

        if not os.path.exists(image_path):
            result['errors'].append(f'文件不存在: {image_path}')
            return result

        # 尝试 pytesseract
        text = self._ocr_with_tesseract(image_path)
        if not text:
            # 回退到 easyocr
            text = self._ocr_with_easyocr(image_path)

        if not text:
            result['errors'].append('OCR识别失败，请检查图片清晰度或手动输入。')
            return result

        result['raw_text'] = text

        # 解析文本提取持仓
        holdings = self._parse_holdings_from_text(text)
        result['holdings'] = holdings
        result['success'] = len(holdings) > 0

        if not holdings:
            result['errors'].append('未能从截图中识别出持仓信息，请确认图片包含基金代码/名称和份额。')

        # 保存到客户仓库
        if client_id and holdings:
            self.save_to_repository(client_id, holdings, source='screenshot',
                                    source_path=image_path)

        return result

    def _ocr_with_tesseract(self, image_path: str) -> str:
        """使用 Tesseract OCR 识别图片"""
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(image_path)
            text = pytesseract.image_to_string(img, lang='chi_sim+eng')
            return text.strip()
        except ImportError:
            return ""
        except Exception:
            return ""

    def _ocr_with_easyocr(self, image_path: str) -> str:
        """使用 EasyOCR 识别图片（备用）"""
        try:
            import easyocr
            reader = easyocr.Reader(['ch_sim', 'en'], gpu=False)
            results = reader.readtext(image_path)
            text = '\n'.join([r[1] for r in results])
            return text.strip()
        except ImportError:
            return ""
        except Exception:
            return ""

    # ==================== Word文档导入 ====================

    def import_from_docx(self, file_path: str, client_id: str = None) -> dict:
        """
        从Word文档导入持仓信息

        Args:
            file_path: .docx 文件路径
            client_id: 客户ID（可选）

        Returns:
            dict: {'success': bool, 'holdings': list, 'errors': list}
        """
        result = {'success': False, 'holdings': [], 'errors': []}

        if not os.path.exists(file_path):
            result['errors'].append(f'文件不存在: {file_path}')
            return result

        try:
            from docx import Document
        except ImportError:
            result['errors'].append('请先安装python-docx: pip install python-docx')
            return result

        try:
            doc = Document(file_path)

            all_text = []
            table_data = []

            # 提取表格
            for table in doc.tables:
                rows_data = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows_data.append(cells)
                table_data.extend(rows_data)

            # 提取段落文本
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    all_text.append(text)

            full_text = '\n'.join(all_text)

            # 优先从表格解析
            if table_data:
                holdings = self._parse_holdings_from_table(table_data)
            else:
                holdings = []

            # 从文本补充
            text_holdings = self._parse_holdings_from_text(full_text)

            # 合并去重
            holdings = self._merge_holdings(holdings, text_holdings)

            result['holdings'] = holdings
            result['success'] = len(holdings) > 0

            if not holdings:
                result['errors'].append('未能从文档中识别出持仓信息，请确保包含基金代码/份额/成本。')

            if client_id and holdings:
                self.save_to_repository(client_id, holdings, source='docx',
                                        source_path=file_path)

        except Exception as e:
            result['errors'].append(f'Word文档解析失败: {str(e)}')

        return result

    # ==================== PDF导入 ====================

    def import_from_pdf(self, file_path: str, client_id: str = None) -> dict:
        """
        从PDF文档导入持仓信息

        Args:
            file_path: .pdf 文件路径
            client_id: 客户ID（可选）

        Returns:
            dict: {'success': bool, 'holdings': list, 'errors': list}
        """
        result = {'success': False, 'holdings': [], 'errors': []}

        if not os.path.exists(file_path):
            result['errors'].append(f'文件不存在: {file_path}')
            return result

        text = self._extract_pdf_text(file_path)
        if not text:
            result['errors'].append('PDF文本提取失败，文件可能为扫描件请使用截图导入。')
            return result

        # 尝试表格解析
        table_data = self._extract_pdf_tables(file_path)

        holdings = []
        if table_data:
            holdings = self._parse_holdings_from_table(table_data)

        # 从文本补充
        text_holdings = self._parse_holdings_from_text(text)
        holdings = self._merge_holdings(holdings, text_holdings)

        result['holdings'] = holdings
        result['success'] = len(holdings) > 0

        if not holdings:
            result['errors'].append('未能从PDF中识别出持仓信息。')

        if client_id and holdings:
            self.save_to_repository(client_id, holdings, source='pdf',
                                    source_path=file_path)

        return result

    def _extract_pdf_text(self, file_path: str) -> str:
        """提取PDF文本"""
        # 优先使用 pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                texts = []
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        texts.append(t)
                return '\n'.join(texts)
        except ImportError:
            pass
        except Exception:
            pass

        # 回退到 PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            texts = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    texts.append(t)
            return '\n'.join(texts)
        except ImportError:
            pass
        except Exception:
            pass

        return ""

    def _extract_pdf_tables(self, file_path: str) -> list:
        """提取PDF表格"""
        try:
            import pdfplumber
            tables = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_tables = page.extract_tables()
                    for t in page_tables:
                        if t:
                            tables.extend(t)
            return tables
        except ImportError:
            return []
        except Exception:
            return []

    # ==================== 浏览器爬取 ====================

    def import_from_url(self, url: str, client_id: str = None,
                        credentials: dict = None, headless: bool = True) -> dict:
        """
        从浏览器链接爬取持仓信息

        支持: 天天基金、支付宝/蚂蚁财富、好买基金等

        Args:
            url: 持仓页面URL
            client_id: 客户ID
            credentials: {'username': ..., 'password': ..., 'platform': ...} 或 None
            headless: 是否无头模式

        Returns:
            dict: {'success': bool, 'holdings': list, 'errors': list, 'screenshot_path': str}
        """
        result = {'success': False, 'holdings': [], 'errors': [], 'screenshot_path': ''}

        platform = self._detect_platform(url)

        if platform == 'eastmoney':
            return self._crawl_eastmoney(url, client_id, credentials)
        elif platform == 'antfortune':
            return self._crawl_antfortune(url, client_id, credentials)

        # 通用爬取: 尝试 Selenium
        return self._crawl_with_selenium(url, client_id, credentials, headless)

    def _detect_platform(self, url: str) -> str:
        """检测平台类型"""
        url_lower = url.lower()
        if 'eastmoney.com' in url_lower or '天天基金' in url_lower:
            return 'eastmoney'
        if 'antfortune.com' in url_lower or '支付宝' in url_lower or '蚂蚁' in url_lower:
            return 'antfortune'
        if 'howbuy.com' in url_lower:
            return 'howbuy'
        if 'lufax.com' in url_lower:
            return 'lufax'
        return 'unknown'

    def _crawl_eastmoney(self, url: str, client_id: str = None,
                         credentials: dict = None) -> dict:
        """爬取天天基金持仓"""
        result = {'success': False, 'holdings': [], 'errors': [], 'screenshot_path': ''}

        try:
            from selenium import webdriver
            from selenium.webdriver.common.by import By
            from selenium.webdriver.support.ui import WebDriverWait
            from selenium.webdriver.support import expected_conditions as EC
            from selenium.webdriver.chrome.options import Options
        except ImportError:
            # 回退到 requests 抓取公开数据
            return self._crawl_eastmoney_api(url, client_id)

        try:
            options = Options()
            options.add_argument('--headless=new')
            options.add_argument('--no-sandbox')
            options.add_argument('--disable-dev-shm-usage')
            options.add_argument('--window-size=1920,1080')

            driver = webdriver.Chrome(options=options)
            driver.get(url)

            wait = WebDriverWait(driver, 15)

            # 如果需要登录
            if credentials and credentials.get('username'):
                try:
                    login_btn = wait.until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, '.login_btn, [class*="login"]'))
                    )
                    login_btn.click()

                    username_input = wait.until(
                        EC.presence_of_element_located((By.CSS_SELECTOR, 'input[name="username"], input[type="text"]'))
                    )
                    username_input.send_keys(credentials['username'])

                    password_input = driver.find_element(By.CSS_SELECTOR, 'input[name="password"], input[type="password"]')
                    password_input.send_keys(credentials.get('password', ''))

                    submit_btn = driver.find_element(By.CSS_SELECTOR, 'button[type="submit"], .login_submit')
                    submit_btn.click()

                    wait.until(lambda d: 'login' not in d.current_url.lower())
                except Exception:
                    result['errors'].append('自动登录失败，请手动登录后提供页面截图。')

            # 截图保存
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            screenshot_path = str(self.clients_dir / (client_id or 'unknown') / 'imports' /
                                  f'screenshot_{timestamp}.png')
            os.makedirs(os.path.dirname(screenshot_path), exist_ok=True)
            driver.save_screenshot(screenshot_path)
            result['screenshot_path'] = screenshot_path

            # 解析页面
            page_text = driver.find_element(By.TAG_NAME, 'body').text
            driver.quit()

            holdings = self._parse_holdings_from_text(page_text)
            result['holdings'] = holdings
            result['success'] = len(holdings) > 0

            if client_id and holdings:
                self.save_to_repository(client_id, holdings, source='url_crawl',
                                        source_path=url)

        except Exception as e:
            result['errors'].append(f'浏览器爬取失败: {str(e)}')

        return result

    def _crawl_eastmoney_api(self, url: str, client_id: str = None) -> dict:
        """通过API抓取天天基金公开数据"""
        result = {'success': False, 'holdings': [], 'errors': [], 'screenshot_path': ''}

        try:
            import requests
            import re

            # 从URL提取基金代码
            codes = re.findall(r'(\d{6})', url)
            if not codes:
                result['errors'].append('未能从URL中提取基金代码')
                return result

            holdings = []
            for code in codes:
                # 天天基金API
                api_url = f'https://fundgz.1234567.com.cn/js/{code}.js'
                resp = requests.get(api_url, headers={
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                    'Referer': 'https://fund.eastmoney.com/'
                }, timeout=10)

                if resp.status_code == 200:
                    match = re.search(r'jsonpgz\((.+)\)', resp.text)
                    if match:
                        data = json.loads(match.group(1))
                        holdings.append({
                            'fund_code': data.get('fundcode', code),
                            'fund_name': data.get('name', ''),
                            'current_nav': float(data.get('dwjz', 0)),
                            'estimated_nav': float(data.get('gsz', 0)),
                            'nav_date': data.get('jzrq', ''),
                            'source': 'eastmoney_api'
                        })

            result['holdings'] = holdings
            result['success'] = len(holdings) > 0

        except Exception as e:
            result['errors'].append(f'API抓取失败: {str(e)}')

        return result

    def _crawl_antfortune(self, url: str, client_id: str = None,
                          credentials: dict = None) -> dict:
        """爬取蚂蚁财富持仓"""
        result = {'success': False, 'holdings': [], 'errors': [], 'screenshot_path': ''}

        result['errors'].append(
            '蚂蚁财富需要扫码登录，建议使用截图导入功能。\n'
            '请截屏您的持仓页面，然后使用 import_from_screenshot 导入。'
        )

        return result

    def _crawl_with_selenium(self, url: str, client_id: str = None,
                             credentials: dict = None, headless: bool = True) -> dict:
        """通用Selenium爬取"""
        result = {'success': False, 'holdings': [], 'errors': [], 'screenshot_path': ''}

        try:
            from selenium import webdriver
            from selenium.webdriver.common.by import By
            from selenium.webdriver.chrome.options import Options
        except ImportError:
            result['errors'].append(
                '需要安装Selenium和ChromeDriver: pip install selenium\n'
                '并将ChromeDriver添加到系统PATH。\n'
                '也可以将页面截图后使用截图导入功能。'
            )
            return result

        try:
            options = Options()
            if headless:
                options.add_argument('--headless=new')
            options.add_argument('--no-sandbox')
            options.add_argument('--disable-dev-shm-usage')
            options.add_argument('--window-size=1920,1080')

            driver = webdriver.Chrome(options=options)
            driver.get(url)

            # 尝试登录
            if credentials and credentials.get('username'):
                self._attempt_login(driver, credentials)

            # 截图
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            client_dir = self.clients_dir / (client_id or 'unknown') / 'imports'
            os.makedirs(str(client_dir), exist_ok=True)
            screenshot_path = str(client_dir / f'screenshot_{timestamp}.png')
            driver.save_screenshot(screenshot_path)
            result['screenshot_path'] = screenshot_path

            # 获取页面文本
            page_text = driver.find_element(By.TAG_NAME, 'body').text
            driver.quit()

            holdings = self._parse_holdings_from_text(page_text)
            result['holdings'] = holdings
            result['success'] = len(holdings) > 0

            if client_id and holdings:
                self.save_to_repository(client_id, holdings, source='url_crawl',
                                        source_path=url)

        except Exception as e:
            result['errors'].append(f'浏览器爬取失败: {str(e)}')

        return result

    def _attempt_login(self, driver, credentials: dict) -> bool:
        """通用登录尝试"""
        try:
            from selenium.webdriver.common.by import By

            username = credentials.get('username', '')
            password = credentials.get('password', '')

            # 查找登录表单
            inputs = driver.find_elements(By.TAG_NAME, 'input')
            text_inputs = [i for i in inputs if i.get_attribute('type') in ('text', 'email', 'tel', None)]
            password_inputs = [i for i in inputs if i.get_attribute('type') == 'password']

            if text_inputs and password_inputs:
                text_inputs[0].send_keys(username)
                password_inputs[0].send_keys(password)

                # 查找提交按钮
                buttons = driver.find_elements(By.TAG_NAME, 'button')
                submit_btns = [b for b in buttons if '登录' in (b.text or '') or 'login' in (b.get_attribute('class') or '').lower()]
                if submit_btns:
                    submit_btns[0].click()
                    return True

            return False
        except Exception:
            return False

    # ==================== 解析引擎 ====================

    def _parse_holdings_from_text(self, text: str) -> list:
        """从文本中解析持仓信息"""
        holdings = []
        if not text:
            return holdings

        # 策略1: 查找表格形式的行
        lines = text.strip().split('\n')
        for line in lines:
            codes = self.FUND_CODE_PATTERN.findall(line)
            if not codes:
                continue

            for code in codes:
                holding = self._extract_holding_from_line(line, code)
                if holding:
                    holdings.append(holding)

        # 策略2: 全文搜索基金代码
        all_codes = self.FUND_CODE_PATTERN.findall(text)
        if len(all_codes) > len(holdings):
            for code in all_codes:
                if not any(h.get('fund_code') == code for h in holdings):
                    # 找到代码附近的上下文
                    idx = text.find(code)
                    context_start = max(0, idx - 100)
                    context_end = min(len(text), idx + 100)
                    context = text[context_start:context_end]
                    holding = self._extract_holding_from_line(context, code)
                    if holding:
                        holdings.append(holding)

        # 标准化
        return [self._normalize_holding(h) for h in holdings]

    def _extract_holding_from_line(self, line: str, fund_code: str) -> Optional[dict]:
        """从一行文本中提取持仓"""
        holding = {'fund_code': fund_code}

        # 提取基金名称
        name_patterns = [
            rf'([^\d\s,，;；\n\r]{{2,20}}?)\s*[\(（]?\s*{fund_code}',
            rf'{fund_code}\s*[\)）]?\s*([^\d\s,，;；\n\r]{{2,20}}?)',
            r'([一-龥A-Za-z]{2,20}?(?:基金|混合|债券|指数|货币|ETF|LOF)[^\d]*)',
        ]
        for pat in name_patterns:
            m = re.search(pat, line)
            if m:
                name = m.group(1).strip()
                # 清理
                name = re.sub(r'[\(\)（）\*]', '', name).strip()
                if len(name) >= 2:
                    holding['fund_name'] = name
                    break

        # 提取份额
        shares_matches = [
            re.search(r'(\d+(?:,\d{3})*(?:\.\d+)?)\s*[份分]', line),
            re.search(r'[份分]\s*[:：]?\s*(\d+(?:,\d{3})*(?:\.\d+)?)', line),
            re.search(r'(?:持有|持仓)\s*(\d+(?:,\d{3})*(?:\.\d+)?)', line),
        ]
        for m in shares_matches:
            if m:
                shares_str = m.group(1).replace(',', '')
                holding['shares'] = float(shares_str)
                break

        # 提取成本/净值
        cost_matches = [
            re.search(r'(?:成本|买入价|净值)[:：\s]*(\d+\.?\d{0,4})', line),
            re.search(r'(\d+\.\d{2,4})\s*(?:元|买入|成本)', line),
        ]
        for m in cost_matches:
            if m:
                holding['cost'] = float(m.group(1))
                break

        # 提取金额
        amount_match = re.search(r'(?:金额|市值|投入)[:：\s]*(\d+(?:,\d{3})*(?:\.\d+)?)', line)
        if amount_match:
            holding['amount'] = float(amount_match.group(1).replace(',', ''))

        # 提取日期
        date_match = self.DATE_PATTERN.search(line)
        if date_match:
            y, m, d = date_match.group(1), date_match.group(2).zfill(2), (date_match.group(3) or '01').zfill(2)
            holding['purchase_date'] = f'{y}-{m}-{d}'

        return holding if len(holding) > 1 else None

    def _parse_holdings_from_table(self, table_data: list) -> list:
        """从表格数据解析持仓

        Args:
            table_data: 二维列表 [[col1, col2, ...], ...]

        Returns:
            holdings列表
        """
        if not table_data or len(table_data) < 2:
            return []

        # 识别表头行
        header_row = None
        col_map = {}

        for i, row in enumerate(table_data):
            row_text = ' '.join(str(c) for c in row).lower()
            if any(kw in row_text for kw in ['代码', '名称', '份额', '成本', '净值', '金额', '日期']):
                header_row = i
                col_map = self._map_table_columns(row)
                break

        if not col_map:
            # 无表头，尝试逐行解析
            holdings = []
            for row in table_data:
                row_text = ' '.join(str(c) for c in row)
                codes = self.FUND_CODE_PATTERN.findall(row_text)
                for code in codes:
                    h = self._extract_holding_from_line(row_text, code)
                    if h:
                        holdings.append(h)
            return holdings

        # 有表头，按列映射解析
        holdings = []
        for row in table_data[header_row + 1:]:
            if not row or all(not c or str(c).strip() == '' for c in row):
                continue

            holding = {}
            for col_idx, field in col_map.items():
                if col_idx < len(row) and row[col_idx]:
                    val = str(row[col_idx]).strip()
                    if field == 'fund_code':
                        match = self.FUND_CODE_PATTERN.search(val)
                        if match:
                            holding['fund_code'] = match.group(1)
                    elif field == 'fund_name':
                        holding['fund_name'] = val
                    elif field == 'shares':
                        try:
                            holding['shares'] = float(val.replace(',', ''))
                        except ValueError:
                            pass
                    elif field == 'cost':
                        try:
                            holding['cost'] = float(val.replace(',', ''))
                        except ValueError:
                            pass
                    elif field == 'amount':
                        try:
                            holding['amount'] = float(val.replace(',', ''))
                        except ValueError:
                            pass
                    elif field == 'purchase_date':
                        holding['purchase_date'] = val

            if holding.get('fund_code'):
                holdings.append(holding)

        return holdings

    def _map_table_columns(self, header_row: list) -> dict:
        """映射表格列名到字段"""
        col_map = {}
        for i, col in enumerate(header_row):
            col_lower = str(col).lower().strip()
            if any(kw in col_lower for kw in ['代码', 'code', '编号']):
                col_map[i] = 'fund_code'
            elif any(kw in col_lower for kw in ['名称', 'name', '基金名', '产品']):
                col_map[i] = 'fund_name'
            elif any(kw in col_lower for kw in ['份额', '份数', 'shares', '持有']):
                col_map[i] = 'shares'
            elif any(kw in col_lower for kw in ['成本', '买入价', 'cost', '净值']):
                col_map[i] = 'cost'
            elif any(kw in col_lower for kw in ['金额', '市值', 'amount', 'value']):
                col_map[i] = 'amount'
            elif any(kw in col_lower for kw in ['日期', 'date', '时间', '买入']):
                col_map[i] = 'purchase_date'
        return col_map

    def _normalize_holding(self, holding: dict) -> dict:
        """标准化持仓记录"""
        normalized = {
            'fund_code': holding.get('fund_code', ''),
            'fund_name': holding.get('fund_name', ''),
            'shares': holding.get('shares'),
            'cost': holding.get('cost'),
            'amount': holding.get('amount'),
            'purchase_date': holding.get('purchase_date'),
            'current_nav': holding.get('current_nav'),
            'source': holding.get('source', 'import'),
            'import_date': datetime.now().strftime('%Y-%m-%d %H:%M'),
        }
        return {k: v for k, v in normalized.items() if v is not None}

    def _merge_holdings(self, existing: list, new: list) -> list:
        """合并去重持仓"""
        seen = set()
        merged = []

        for h in existing + new:
            key = h.get('fund_code', '')
            if key and key not in seen:
                seen.add(key)
                merged.append(h)

        return merged

    def validate_holdings(self, holdings: list) -> tuple:
        """验证持仓数据

        Returns:
            (valid_holdings, errors)
        """
        valid = []
        errors = []

        for i, h in enumerate(holdings):
            row_errors = []

            fund_code = h.get('fund_code', '')
            if not fund_code or not re.match(r'^\d{6}$', fund_code):
                row_errors.append(f'第{i+1}行: 基金代码无效 ({fund_code})')

            if 'shares' in h and h['shares'] is not None and h['shares'] <= 0:
                row_errors.append(f'第{i+1}行: 份额必须大于0')

            if 'cost' in h and h['cost'] is not None and h['cost'] <= 0:
                row_errors.append(f'第{i+1}行: 成本必须大于0')

            if row_errors:
                errors.extend(row_errors)
            else:
                valid.append(h)

        return valid, errors

    # ==================== 导出功能 ====================

    def export_to_excel(self, holdings: list, output_path: str = None,
                        client_id: str = None) -> str:
        """
        导出持仓到Excel

        Args:
            holdings: 持仓列表
            output_path: 输出路径（可选，默认自动生成）
            client_id: 客户ID

        Returns:
            str: 输出文件路径
        """
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            raise ImportError('请先安装openpyxl: pip install openpyxl')

        # 生成输出路径
        if not output_path:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_dir = self.clients_dir / (client_id or 'export') / 'reports'
            os.makedirs(str(output_dir), exist_ok=True)
            output_path = str(output_dir / f'holdings_{timestamp}.xlsx')

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = '持仓明细'

        # 样式
        header_font = Font(name='微软雅黑', bold=True, size=11, color='FFFFFF')
        header_fill = PatternFill(start_color='2F5496', end_color='2F5496', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center')
        cell_alignment = Alignment(horizontal='center', vertical='center')
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        # 计算额外列
        has_amount = any('amount' in h for h in holdings)
        has_date = any('purchase_date' in h for h in holdings)

        # 表头
        headers = ['序号', '基金代码', '基金名称', '持有份额', '成本价(元)', '当前净值(元)']
        if has_amount:
            headers.append('投入金额(元)')
        headers.extend(['持仓市值(元)', '盈亏(元)', '收益率(%)'])
        if has_date:
            headers.append('买入日期')
        headers.append('数据来源')

        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border

        # 数据行
        red_font = Font(color='FF0000')
        green_font = Font(color='008000')

        for i, h in enumerate(holdings):
            row = i + 2
            shares = h.get('shares', 0) or 0
            cost = h.get('cost', 0) or 0
            nav = h.get('current_nav', cost)
            amount = h.get('amount', shares * cost)

            market_value = shares * nav
            profit = market_value - amount
            profit_pct = (profit / amount * 100) if amount > 0 else 0

            data = [
                i + 1,
                h.get('fund_code', ''),
                h.get('fund_name', ''),
                shares,
                round(cost, 4),
                round(nav, 4),
            ]
            if has_amount:
                data.append(round(amount, 2))
            data.extend([
                round(market_value, 2),
                round(profit, 2),
                round(profit_pct, 2),
            ])
            if has_date:
                data.append(h.get('purchase_date', ''))
            data.append(h.get('source', '导入'))

            for col, val in enumerate(data, 1):
                cell = ws.cell(row=row, column=col, value=val)
                cell.alignment = cell_alignment
                cell.border = thin_border

            # 盈亏颜色
            profit_col = 7 + (1 if has_amount else 0)
            profit_cell = ws.cell(row=row, column=profit_col)
            pct_cell = ws.cell(row=row, column=profit_col + 1)
            if profit > 0:
                profit_cell.font = red_font
                pct_cell.font = red_font
            elif profit < 0:
                profit_cell.font = green_font
                pct_cell.font = green_font

        # 汇总行
        summary_row = len(holdings) + 2
        ws.cell(row=summary_row, column=1, value='合计').font = Font(bold=True)

        total_shares_col = 4
        total_shares = sum(h.get('shares', 0) or 0 for h in holdings)
        ws.cell(row=summary_row, column=total_shares_col, value=total_shares)

        total_amount = sum(
            h.get('amount', (h.get('shares', 0) or 0) * (h.get('cost', 0) or 0))
            for h in holdings
        )
        total_value = sum(
            (h.get('shares', 0) or 0) * (h.get('current_nav', h.get('cost', 0) or 0))
            for h in holdings
        )
        total_profit = total_value - total_amount
        total_profit_pct = (total_profit / total_amount * 100) if total_amount > 0 else 0

        value_col_start = 6
        if has_amount:
            ws.cell(row=summary_row, column=value_col_start, value=round(total_amount, 2))
            value_col_start += 1
        ws.cell(row=summary_row, column=value_col_start, value=round(total_value, 2))
        ws.cell(row=summary_row, column=value_col_start + 1, value=round(total_profit, 2))
        ws.cell(row=summary_row, column=value_col_start + 2, value=round(total_profit_pct, 2))

        # 调整列宽
        col_widths = [6, 12, 25, 14, 14, 14]
        if has_amount:
            col_widths.append(16)
        col_widths.extend([16, 14, 12])
        if has_date:
            col_widths.append(14)
        col_widths.append(12)

        for col, width in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(col)].width = width

        # 冻结表头
        ws.freeze_panes = 'A2'

        # 添加汇总sheet
        ws2 = wb.create_sheet('持仓汇总')
        summary_data = [
            ['项目', '数值'],
            ['持仓数量', f'{len(holdings)}只'],
            ['总投入金额', f'¥{total_amount:,.2f}'],
            ['当前总市值', f'¥{total_value:,.2f}'],
            ['总盈亏', f'¥{total_profit:,.2f}'],
            ['总收益率', f'{total_profit_pct:+.2f}%'],
            ['导出时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
        ]
        for r, (label, val) in enumerate(summary_data, 1):
            ws2.cell(row=r, column=1, value=label).font = Font(bold=True, name='微软雅黑')
            ws2.cell(row=r, column=2, value=val).font = Font(name='微软雅黑')
            if r == 5 and total_profit > 0:
                ws2.cell(row=r, column=2).font = Font(name='微软雅黑', color='FF0000')
            elif r == 5 and total_profit < 0:
                ws2.cell(row=r, column=2).font = Font(name='微软雅黑', color='008000')

        ws2.column_dimensions['A'].width = 18
        ws2.column_dimensions['B'].width = 25

        wb.save(output_path)
        return output_path

    def export_to_csv(self, holdings: list, output_path: str = None,
                      client_id: str = None) -> str:
        """导出持仓到CSV"""
        import csv

        if not output_path:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_dir = self.clients_dir / (client_id or 'export') / 'reports'
            os.makedirs(str(output_dir), exist_ok=True)
            output_path = str(output_dir / f'holdings_{timestamp}.csv')

        fieldnames = [
            'fund_code', 'fund_name', 'shares', 'cost', 'current_nav',
            'amount', 'market_value', 'profit', 'profit_pct',
            'purchase_date', 'source', 'import_date'
        ]

        with open(output_path, 'w', newline='', encoding='utf-8-sig') as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction='ignore')
            writer.writeheader()

            for h in holdings:
                shares = h.get('shares', 0) or 0
                cost = h.get('cost', 0) or 0
                nav = h.get('current_nav', cost)
                amount = h.get('amount', shares * cost)
                market_value = shares * nav

                row = {
                    'fund_code': h.get('fund_code', ''),
                    'fund_name': h.get('fund_name', ''),
                    'shares': shares,
                    'cost': round(cost, 4),
                    'current_nav': round(nav, 4),
                    'amount': round(amount, 2),
                    'market_value': round(market_value, 2),
                    'profit': round(market_value - amount, 2),
                    'profit_pct': round((market_value - amount) / amount * 100, 2) if amount > 0 else 0,
                    'purchase_date': h.get('purchase_date', ''),
                    'source': h.get('source', 'import'),
                    'import_date': h.get('import_date', datetime.now().strftime('%Y-%m-%d')),
                }
                writer.writerow(row)

        return output_path

    # ==================== 客户持仓仓库 ====================

    def save_to_repository(self, client_id: str, holdings: list,
                           source: str = 'manual', source_path: str = None) -> str:
        """
        保存持仓到客户仓库

        Args:
            client_id: 客户ID
            holdings: 持仓列表
            source: 数据来源 (screenshot/docx/pdf/url_crawl/manual)
            source_path: 源文件路径

        Returns:
            str: 保存路径
        """
        client_dir = self.clients_dir / client_id
        client_dir.mkdir(parents=True, exist_ok=True)
        (client_dir / 'imports').mkdir(exist_ok=True)
        (client_dir / 'reports').mkdir(exist_ok=True)
        (client_dir / 'history').mkdir(exist_ok=True)

        # 添加元数据
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        for h in holdings:
            h.setdefault('source', source)
            h.setdefault('import_date', datetime.now().strftime('%Y-%m-%d %H:%M'))
            h['import_batch'] = timestamp

        # 读取现有持仓
        holdings_file = client_dir / 'holdings.json'
        existing = {}
        if holdings_file.exists():
            try:
                with open(holdings_file, 'r', encoding='utf-8') as f:
                    existing_list = json.load(f)
                    for h in existing_list:
                        code = h.get('fund_code', '')
                        if code:
                            existing[code] = h
            except Exception:
                pass

        # 合并：新数据覆盖旧数据同代码的持仓
        for h in holdings:
            code = h.get('fund_code', '')
            if code:
                existing[code] = h

        merged = list(existing.values())

        # 保存
        with open(holdings_file, 'w', encoding='utf-8') as f:
            json.dump(merged, f, ensure_ascii=False, indent=2)

        # 保存历史快照
        history_file = client_dir / 'history' / f'holdings_{timestamp}.json'
        with open(history_file, 'w', encoding='utf-8') as f:
            json.dump(merged, f, ensure_ascii=False, indent=2)

        # 保存导入记录
        import_log = client_dir / 'imports' / f'import_{timestamp}.json'
        import_record = {
            'timestamp': timestamp,
            'source': source,
            'source_path': source_path,
            'count': len(holdings),
            'holdings': holdings
        }
        with open(import_log, 'w', encoding='utf-8') as f:
            json.dump(import_record, f, ensure_ascii=False, indent=2)

        # 更新仓库索引
        self._update_repository_index(client_id, merged)

        return str(holdings_file)

    def load_client_holdings(self, client_id: str) -> list:
        """加载客户当前持仓"""
        holdings_file = self.clients_dir / client_id / 'holdings.json'
        if not holdings_file.exists():
            return []
        try:
            with open(holdings_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return []

    def get_import_history(self, client_id: str) -> list:
        """获取客户导入历史"""
        imports_dir = self.clients_dir / client_id / 'imports'
        if not imports_dir.exists():
            return []

        history = []
        for f in sorted(imports_dir.glob('import_*.json'), reverse=True):
            try:
                with open(f, 'r', encoding='utf-8') as fh:
                    history.append(json.load(fh))
            except Exception:
                pass
        return history

    def get_history_snapshots(self, client_id: str) -> list:
        """获取客户持仓历史快照"""
        history_dir = self.clients_dir / client_id / 'history'
        if not history_dir.exists():
            return []

        snapshots = []
        for f in sorted(history_dir.glob('holdings_*.json'), reverse=True):
            snapshots.append({
                'timestamp': f.stem.replace('holdings_', ''),
                'path': str(f)
            })
        return snapshots

    def _update_repository_index(self, client_id: str, holdings: list):
        """更新仓库索引"""
        index_file = self.clients_dir / '_index.json'
        index = {}
        if index_file.exists():
            try:
                with open(index_file, 'r', encoding='utf-8') as f:
                    index = json.load(f)
            except Exception:
                pass

        # 计算市值
        total_value = 0
        for h in holdings:
            shares = h.get('shares', 0) or 0
            nav = h.get('current_nav', h.get('cost', 0) or 0)
            total_value += shares * nav

        index[client_id] = {
            'name': index.get(client_id, {}).get('name', client_id),
            'holdings_count': len(holdings),
            'total_value': round(total_value, 2),
            'last_updated': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'fund_codes': [h.get('fund_code') for h in holdings]
        }

        with open(index_file, 'w', encoding='utf-8') as f:
            json.dump(index, f, ensure_ascii=False, indent=2)

    def list_clients(self) -> dict:
        """列出所有客户"""
        index_file = self.clients_dir / '_index.json'
        if index_file.exists():
            try:
                with open(index_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception:
                pass
        return {}

    def delete_client(self, client_id: str) -> bool:
        """删除客户及所有数据"""
        import shutil
        client_dir = self.clients_dir / client_id
        if client_dir.exists():
            shutil.rmtree(str(client_dir))
            self._update_repository_index(client_id, [])
            return True
        return False


def main():
    """测试导入引擎"""
    importer = HoldingsImporter()
    print("=" * 60)
    print("  持仓导入引擎 v1.0 测试")
    print("=" * 60)

    # 测试文本解析
    test_text = """
    基金持仓明细
    001924 华夏国企改革混合 10000份 成本1.500 2025-01-15
    110022 易方达消费行业 5000份 成本3.000 2025-03-20
    006912 长城久泰沪深300指数C 14000份 成本1.000 2025-06-01
    """

    print("\n【测试文本解析】")
    print(test_text)
    holdings = importer._parse_holdings_from_text(test_text)
    for h in holdings:
        print(f"  {h.get('fund_code')} {h.get('fund_name', '?')} "
              f"份额:{h.get('shares')} 成本:{h.get('cost')} 日期:{h.get('purchase_date', '')}")

    # 测试表格解析
    test_table = [
        ['基金代码', '基金名称', '持有份额', '成本价'],
        ['001924', '华夏国企改革混合', '10000', '1.500'],
        ['110022', '易方达消费行业', '5000', '3.000'],
        ['006912', '长城久泰沪深300指数C', '14000', '1.000'],
    ]

    print("\n【测试表格解析】")
    table_holdings = importer._parse_holdings_from_table(test_table)
    for h in table_holdings:
        print(f"  {h.get('fund_code')} {h.get('fund_name', '?')} "
              f"份额:{h.get('shares')} 成本:{h.get('cost')}")

    # 测试保存到仓库
    print("\n【测试仓库保存】")
    path = importer.save_to_repository('test_client', holdings, source='test')
    print(f"  保存到: {path}")

    # 测试加载
    loaded = importer.load_client_holdings('test_client')
    print(f"  加载到 {len(loaded)} 条持仓")

    # 测试导出
    if holdings:
        try:
            excel_path = importer.export_to_excel(holdings, client_id='test_client')
            print(f"\n【Excel导出】 {excel_path}")
        except ImportError:
            print("\n【Excel导出】 跳过 (需安装openpyxl)")

        csv_path = importer.export_to_csv(holdings, client_id='test_client')
        print(f"【CSV导出】 {csv_path}")

    print("\n" + "=" * 60)
    print("  测试完成")
    print("=" * 60)


if __name__ == '__main__':
    main()
